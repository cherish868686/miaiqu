"""
Word文档解析器 (.docx)
提取文本段落和表格内容
"""
from docx import Document


def parse_docx(filepath):
    """解析Word文档，返回文本和表格内容"""
    doc = Document(filepath)
    texts = []
    tables = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            texts.append({
                'index': i,
                'content': text,
                'style': para.style.name if para.style else None
            })

    for t_idx, table in enumerate(doc.tables):
        table_data = {
            'index': t_idx,
            'headers': [],
            'rows': []
        }
        for r_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            if r_idx == 0:
                table_data['headers'] = row_data
            else:
                table_data['rows'].append({
                    'row_index': r_idx,
                    'cells': row_data
                })
        tables.append(table_data)

    return {'texts': texts, 'tables': tables}
