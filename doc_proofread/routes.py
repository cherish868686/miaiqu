"""
文档校对API路由
注册到密信本Flask应用中
"""
import os
import json
import uuid
import difflib
from datetime import datetime, timedelta
from flask import request, jsonify
from werkzeug.utils import secure_filename

# 延迟导入，避免缺少可选依赖时整个应用无法启动
_parse_docx = None
_parse_xlsx = None
_parse_pdf = None
_parse_pptx = None
_parse_image = None
_compare_texts = None
_compare_tables = None


def _lazy_imports():
    global _parse_docx, _parse_xlsx, _parse_pdf, _parse_pptx, _parse_image
    global _compare_texts, _compare_tables
    if _parse_docx is not None:
        return
    from doc_proofread.parsers.docx_parser import parse_docx
    from doc_proofread.parsers.xlsx_parser import parse_xlsx
    from doc_proofread.parsers.pdf_parser import parse_pdf
    from doc_proofread.parsers.pptx_parser import parse_pptx
    from doc_proofread.parsers.image_parser import parse_image
    from doc_proofread.comparator.text_comparator import compare_texts
    from doc_proofread.comparator.table_comparator import compare_tables
    _parse_docx = parse_docx
    _parse_xlsx = parse_xlsx
    _parse_pdf = parse_pdf
    _parse_pptx = parse_pptx
    _parse_image = parse_image
    _compare_texts = compare_texts
    _compare_tables = compare_tables

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'data', 'proofread_uploads')
HISTORY_FOLDER = os.path.join(BASE_DIR, 'data', 'proofread_history')
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'xlsx', 'xls', 'pptx', 'ppt', 'jpg', 'jpeg', 'png'}
MAX_CONTENT_LENGTH = 50 * 1024 * 1024  # 50MB


def allowed_file(filename):
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    return ext in ALLOWED_EXTENSIONS


def get_file_type(filename):
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    type_map = {
        'pdf': 'pdf',
        'docx': 'word', 'doc': 'word',
        'xlsx': 'excel', 'xls': 'excel',
        'pptx': 'ppt', 'ppt': 'ppt',
        'jpg': 'image', 'jpeg': 'image', 'png': 'image',
    }
    return type_map.get(ext, 'unknown')


def parse_file(filepath, filename):
    _lazy_imports()
    file_type = get_file_type(filename)
    if file_type == 'word':
        return _parse_docx(filepath)
    elif file_type == 'excel':
        return _parse_xlsx(filepath)
    elif file_type == 'pdf':
        return _parse_pdf(filepath)
    elif file_type == 'ppt':
        return _parse_pptx(filepath)
    elif file_type == 'image':
        return _parse_image(filepath)
    else:
        raise ValueError(f"不支持的文件类型: {filename}")


def cleanup_old_history():
    """清理超过一周的历史记录"""
    now = datetime.now()
    if not os.path.exists(HISTORY_FOLDER):
        return
    for filename in os.listdir(HISTORY_FOLDER):
        filepath = os.path.join(HISTORY_FOLDER, filename)
        if os.path.isfile(filepath) and filename.endswith('.json'):
            try:
                file_time = datetime.fromtimestamp(os.path.getmtime(filepath))
                if now - file_time > timedelta(days=7):
                    os.remove(filepath)
            except Exception:
                pass


def perform_comparison(parsed_docs):
    _lazy_imports()
    results = []
    for i in range(len(parsed_docs)):
        for j in range(i + 1, len(parsed_docs)):
            doc_a = parsed_docs[i]
            doc_b = parsed_docs[j]

            comparison = {
                'doc_a': doc_a['name'],
                'doc_b': doc_b['name'],
                'type_a': doc_a['file_type'],
                'type_b': doc_b['file_type'],
                'text_diff': None,
                'table_diff': None,
                'summary': {}
            }

            content_a = doc_a['content']
            content_b = doc_b['content']

            texts_a = content_a.get('texts', [])
            texts_b = content_b.get('texts', [])
            if texts_a or texts_b:
                comparison['text_diff'] = _compare_texts(texts_a, texts_b)

            tables_a = content_a.get('tables', [])
            tables_b = content_b.get('tables', [])
            if tables_a or tables_b:
                comparison['table_diff'] = _compare_tables(tables_a, tables_b)

            summary = generate_summary(comparison)
            comparison['summary'] = summary
            results.append(comparison)

    return results


def generate_summary(comparison):
    summary = {
        'total_text_differences': 0,
        'total_table_differences': 0,
        'total_table_rows': 0,
        'consistent_table_rows': 0,
        'overall_status': '一致'
    }

    if comparison.get('text_diff'):
        text_diff = comparison['text_diff']
        summary['total_text_differences'] = text_diff.get('difference_count', 0)
        if text_diff.get('difference_count', 0) > 0:
            summary['overall_status'] = '不一致'

    if comparison.get('table_diff'):
        table_diff = comparison['table_diff']
        total_rows = 0
        diff_rows = 0
        for table in table_diff.get('tables', []):
            rows = table.get('rows', [])
            total_rows += len(rows)
            diff_rows += sum(1 for r in rows if r.get('status') == '不一致')

        summary['total_table_rows'] = total_rows
        summary['consistent_table_rows'] = total_rows - diff_rows
        summary['total_table_differences'] = diff_rows
        if diff_rows > 0:
            summary['overall_status'] = '不一致'

    return summary


def determine_overall_status(results):
    for r in results:
        if r.get('summary', {}).get('overall_status') == '不一致':
            return '不一致'
    return '一致'


def generate_annotated_docx(original_path, record):
    """生成带标注的Word文档"""
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_COLOR_INDEX
    import shutil

    # 复制原文件
    annotated_name = f'annotated_{uuid.uuid4().hex}.docx'
    annotated_path = os.path.join(UPLOAD_FOLDER, annotated_name)
    shutil.copy2(original_path, annotated_path)

    doc = Document(annotated_path)
    results = record.get('results', [])

    # 收集所有文本差异的行索引和内容
    diff_lines = {}  # {line_index: diff_info}
    for result in results:
        text_diff = result.get('text_diff')
        if not text_diff:
            continue
        for diff in text_diff.get('differences', []):
            for detail in diff.get('details', []):
                line_idx = detail.get('line_a_index')
                if line_idx is not None:
                    diff_lines[line_idx] = {
                        'type': diff['type'],
                        'content_a': detail.get('content_a', ''),
                        'content_b': detail.get('content_b', ''),
                        'char_diff': detail.get('char_diff')
                    }

    # 在文档段落中标注差异
    para_idx = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if para_idx in diff_lines:
            diff_info = diff_lines[para_idx]
            # 清空段落
            for run in para.runs:
                run.text = ''

            if not para.runs:
                # 没有runs，添加新的
                run = para.add_run(text)
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                # 在第一个run中设置标注文本
                if diff_info['type'] == 'replace':
                    para.runs[0].text = diff_info['content_a']
                    para.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
                    # 添加校对说明
                    note_run = para.add_run(' → 校对为：' + diff_info['content_b'])
                    note_run.font.color.rgb = RGBColor(0, 128, 0)
                    note_run.font.size = Pt(9)
                    note_run.font.italic = True
                elif diff_info['type'] == 'delete':
                    para.runs[0].text = diff_info['content_a']
                    para.runs[0].font.highlight_color = WD_COLOR_INDEX.RED
                    para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                    note_run = para.add_run(' [已删除]')
                    note_run.font.color.rgb = RGBColor(255, 0, 0)
                    note_run.font.size = Pt(9)
                    note_run.font.italic = True
                elif diff_info['type'] == 'insert':
                    para.runs[0].text = diff_info.get('content_b', '')
                    para.runs[0].font.highlight_color = WD_COLOR_INDEX.GREEN
                    note_run = para.add_run(' [新增]')
                    note_run.font.color.rgb = RGBColor(0, 128, 0)
                    note_run.font.size = Pt(9)
                    note_run.font.italic = True
                else:
                    para.runs[0].text = text
                    para.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW

        para_idx += 1

    # 在文档末尾添加校对摘要
    doc.add_paragraph('')
    summary_para = doc.add_paragraph('—— 校对摘要 ——')
    summary_para.alignment = 1  # center
    summary_run = summary_para.runs[0] if summary_para.runs else summary_para.add_run('—— 校对摘要 ——')
    summary_run.font.bold = True
    summary_run.font.size = Pt(12)

    for result in results:
        s = result.get('summary', {})
        doc.add_paragraph(f"文件: {result.get('doc_a', '')} vs {result.get('doc_b', '')}")
        doc.add_paragraph(f"文本差异: {s.get('total_text_differences', 0)} 处")
        doc.add_paragraph(f"表格差异: {s.get('total_table_differences', 0)} 行")
        doc.add_paragraph(f"整体状态: {s.get('overall_status', '未知')}")

    doc.save(annotated_path)
    return annotated_path


def register_proofread_routes(app, logger):
    """注册文档校对API路由到Flask应用"""

    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(HISTORY_FOLDER, exist_ok=True)

    @app.route('/api/proofread/upload', methods=['POST'])
    def proofread_upload():
        """上传文件接口 - 支持role参数区分原文本/校对文本"""
        cleanup_old_history()

        if 'files' not in request.files:
            return jsonify({'status': 'error', 'message': '没有上传文件'}), 400

        files = request.files.getlist('files')
        if not files or all(f.filename == '' for f in files):
            return jsonify({'status': 'error', 'message': '没有选择文件'}), 400

        role = request.form.get('role', '')  # original 或 proofread

        uploaded = []
        for f in files:
            if f.filename == '':
                continue
            if not allowed_file(f.filename):
                return jsonify({'status': 'error', 'message': f'不支持的文件类型: {f.filename}'}), 400

            original_name = secure_filename(f.filename)
            ext = original_name.rsplit('.', 1)[-1] if '.' in original_name else ''
            unique_name = f"{uuid.uuid4().hex}.{ext}" if ext else uuid.uuid4().hex
            save_path = os.path.join(UPLOAD_FOLDER, unique_name)
            f.save(save_path)

            uploaded.append({
                'id': unique_name,
                'original_name': original_name,
                'file_type': get_file_type(original_name),
                'path': save_path,
                'role': role
            })

        if not uploaded:
            return jsonify({'status': 'error', 'message': '没有有效的文件被上传'}), 400

        role_text = '原文本' if role == 'original' else '校对文本' if role == 'proofread' else '文件'
        logger.info(f"文档校对: 上传了 {len(uploaded)} 个{role_text}")
        return jsonify({
            'status': 'success',
            'message': f'成功上传{role_text}',
            'files': uploaded
        })

    @app.route('/api/proofread/compare', methods=['POST'])
    def proofread_compare():
        """比对文档接口 - 区分原文本和校对文本"""
        cleanup_old_history()

        data = request.get_json(force=True, silent=True) or {}

        # 支持新格式(original/proofread)和旧格式(files)
        original_info = data.get('original')
        proofread_info = data.get('proofread')
        original_file_id = None

        if original_info and proofread_info:
            # 新格式：区分原文本和校对文本
            files_info = [original_info, proofread_info]
            original_file_id = original_info.get('id')
        elif 'files' in data:
            # 旧格式兼容
            files_info = data['files']
            if len(files_info) < 2:
                return jsonify({'status': 'error', 'message': '至少需要两个文件进行比对'}), 400
            original_file_id = files_info[0].get('id')
        else:
            return jsonify({'status': 'error', 'message': '请提供原文本和校对文本'}), 400

        parsed_docs = []
        for fi in files_info:
            filepath = os.path.join(UPLOAD_FOLDER, fi['id'])
            if not os.path.exists(filepath):
                return jsonify({'status': 'error', 'message': f"文件不存在: {fi['original_name']}"}), 400
            try:
                content = parse_file(filepath, fi['original_name'])
                parsed_docs.append({
                    'name': fi['original_name'],
                    'file_type': fi.get('file_type', get_file_type(fi['original_name'])),
                    'content': content
                })
            except Exception as e:
                logger.error(f"解析文件 {fi['original_name']} 失败: {str(e)}")
                return jsonify({'status': 'error', 'message': f"解析文件 {fi['original_name']} 失败: {str(e)}"}), 500

        try:
            comparison_results = perform_comparison(parsed_docs)
        except Exception as e:
            logger.error(f"比对失败: {str(e)}")
            return jsonify({'status': 'error', 'message': f'比对失败: {str(e)}'}), 500

        # 保存历史记录
        history_id = uuid.uuid4().hex
        history_record = {
            'id': history_id,
            'timestamp': datetime.now().isoformat(),
            'files': [d['name'] for d in parsed_docs],
            'file_types': [d['file_type'] for d in parsed_docs],
            'results': comparison_results,
            'original_file_id': original_file_id,
            'original_file_name': files_info[0].get('original_name', '')
        }

        history_path = os.path.join(HISTORY_FOLDER, f'{history_id}.json')
        with open(history_path, 'w', encoding='utf-8') as hf:
            json.dump(history_record, hf, ensure_ascii=False, indent=2)

        logger.info(f"文档校对: 比对完成, history_id={history_id}")
        return jsonify({
            'status': 'success',
            'message': '比对完成',
            'history_id': history_id,
            'original_file_id': original_file_id,
            'results': comparison_results
        })

    @app.route('/api/proofread/download-annotated/<history_id>', methods=['GET'])
    def proofread_download_annotated(history_id):
        """下载标注后的原文档"""
        history_path = os.path.join(HISTORY_FOLDER, f'{history_id}.json')
        if not os.path.exists(history_path):
            return jsonify({'status': 'error', 'message': '历史记录不存在'}), 404

        try:
            with open(history_path, 'r', encoding='utf-8') as f:
                record = json.load(f)

            original_file_id = record.get('original_file_id')
            original_file_name = record.get('original_file_name', 'document')

            if not original_file_id:
                return jsonify({'status': 'error', 'message': '原文件信息缺失'}), 400

            original_path = os.path.join(UPLOAD_FOLDER, original_file_id)
            if not os.path.exists(original_path):
                return jsonify({'status': 'error', 'message': '原文件已过期'}), 404

            file_type = get_file_type(original_file_name)

            # 对于Word文档，生成带标注的版本
            if file_type == 'word':
                annotated_path = generate_annotated_docx(original_path, record)
                from flask import send_file
                base_name = original_file_name.rsplit('.', 1)[0] if '.' in original_file_name else original_file_name
                download_name = f'{base_name}_校对标注.docx'
                return send_file(annotated_path, as_attachment=True, download_name=download_name)
            else:
                # 其他格式直接下载原文件
                from flask import send_file
                return send_file(original_path, as_attachment=True, download_name=original_file_name)

        except Exception as e:
            logger.error(f"下载标注文档失败: {str(e)}")
            return jsonify({'status': 'error', 'message': f'下载失败: {str(e)}'}), 500

    @app.route('/api/proofread/history', methods=['GET'])
    def proofread_history():
        """获取历史记录列表"""
        cleanup_old_history()
        history_list = []

        if os.path.exists(HISTORY_FOLDER):
            for filename in os.listdir(HISTORY_FOLDER):
                if filename.endswith('.json'):
                    filepath = os.path.join(HISTORY_FOLDER, filename)
                    try:
                        with open(filepath, 'r', encoding='utf-8') as f:
                            record = json.load(f)
                        history_list.append({
                            'id': record.get('id'),
                            'timestamp': record.get('timestamp'),
                            'files': record.get('files', []),
                            'file_types': record.get('file_types', []),
                            'overall_status': determine_overall_status(record.get('results', []))
                        })
                    except Exception:
                        pass

        history_list.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
        return jsonify({'status': 'success', 'data': history_list})

    @app.route('/api/proofread/history/<history_id>', methods=['GET'])
    def proofread_history_detail(history_id):
        """获取历史记录详情"""
        filepath = os.path.join(HISTORY_FOLDER, f'{history_id}.json')
        if not os.path.exists(filepath):
            return jsonify({'status': 'error', 'message': '历史记录不存在'}), 404

        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                record = json.load(f)
            return jsonify({'status': 'success', 'data': record})
        except Exception as e:
            return jsonify({'status': 'error', 'message': f'读取历史记录失败: {str(e)}'}), 500

    @app.route('/api/proofread/history/<history_id>', methods=['DELETE'])
    def proofread_history_delete(history_id):
        """删除历史记录"""
        filepath = os.path.join(HISTORY_FOLDER, f'{history_id}.json')
        if not os.path.exists(filepath):
            return jsonify({'status': 'error', 'message': '历史记录不存在'}), 404

        try:
            os.remove(filepath)
            logger.info(f"文档校对: 删除历史记录 {history_id}")
            return jsonify({'status': 'success', 'message': '历史记录已删除'})
        except Exception as e:
            return jsonify({'status': 'error', 'message': f'删除失败: {str(e)}'}), 500
