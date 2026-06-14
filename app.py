"""
新城小米虾 - 集成密信本智能信息监控系统
整合了密信本的信息监控、文档校对等核心能力
"""
from flask import Flask, render_template, request, jsonify, send_from_directory, Response, redirect, url_for
import logging
import os
import threading
import csv
import io
import re
from datetime import datetime
from config import Config
from database import Database
from api_routes import register_api_routes
from doc_proofread.routes import register_proofread_routes
from resource_pool_routes import rp_bp
from parts_price_routes import pp_bp
from cost_library_routes import cl_bp

# 配置基础路径
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

app = Flask(__name__,
           static_folder=os.path.join(BASE_DIR, 'static'),
           template_folder=os.path.join(BASE_DIR, 'templates'))
app.config['TEMPLATES_AUTO_RELOAD'] = True

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    filename=os.path.join(BASE_DIR, 'app.log')
)
logger = logging.getLogger(__name__)

# 全局配置和数据库实例
app_config = Config()
app_config.load_custom()
db = Database(app_config.database_path)

# 注册API路由
register_api_routes(app, app_config, db, logger)

# 注册文档校对路由
register_proofread_routes(app, logger)

# 注册资源池配置路由
app.register_blueprint(rp_bp)

# 注册配件时价路由
app.register_blueprint(pp_bp)

# 注册成本库路由
app.register_blueprint(cl_bp)

# 调度器状态
scheduler_status = {'running': False, 'started_at': None}


@app.route('/')
def index():
    """首页 - 新城小米虾主界面"""
    try:
        logger.info("访问首页")
        return render_template('index_new.html')
    except Exception as e:
        logger.exception("首页渲染异常")
        return f"服务器错误: {str(e)}", 500


@app.route('/mixin')
def mixin_page():
    """密信本页面 - 信息监控中心"""
    return redirect(url_for('index'))


@app.route('/static/<path:filename>')
def static_files(filename):
    try:
        return send_from_directory(app.static_folder, filename)
    except Exception as e:
        logger.error(f"静态文件访问错误: {str(e)}")
        return "资源不存在", 404


@app.route('/send-test-email', methods=['POST'])
def send_test_email():
    """发送测试邮件"""
    try:
        from email_sender import EmailSender
        sender = EmailSender(app_config)
        sender.send_test_email()
        db.save_task_log('test_email', 'success', '测试邮件发送成功')
        logger.info("测试邮件发送成功")
        return jsonify({"status": "success", "message": "测试邮件发送成功！"})
    except Exception as e:
        db.save_task_log('test_email', 'error', str(e))
        logger.exception("测试邮件发送失败")
        return jsonify({"status": "error", "message": str(e)})


@app.route('/start-scheduler', methods=['POST'])
def start_scheduler():
    """启动定时任务，支持自定义配置"""
    global scheduler_status
    try:
        if scheduler_status['running']:
            return jsonify({"status": "info", "message": "定时任务已在运行中"})

        from crawlers import OperatorCrawler, MarketCrawler
        from email_sender import EmailSender
        from scheduler import schedule_jobs

        schedule_config = request.get_json(silent=True) or {}

        operator_crawler = OperatorCrawler(app_config)
        market_crawler = MarketCrawler(app_config)
        email_sender = EmailSender(app_config)

        def run_scheduler():
            schedule_jobs(operator_crawler, market_crawler, email_sender, db, schedule_config)

        scheduler_thread = threading.Thread(target=run_scheduler, daemon=True)
        scheduler_thread.start()
        scheduler_status = {
            'running': True,
            'started_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'config': schedule_config
        }
        db.save_task_log('scheduler', 'success', '定时任务已启动')
        logger.info(f"定时任务已启动，配置: {schedule_config}")

        op_hour = schedule_config.get('operator_hour', 9)
        op_minute = schedule_config.get('operator_minute', 0)
        mkt_hour = schedule_config.get('market_hour', 17)
        mkt_minute = schedule_config.get('market_minute', 0)
        op_range = schedule_config.get('operator_range', 'all')
        mkt_range = schedule_config.get('market_range', 'all')
        range_names = {'all': '爬取+发送邮件', 'crawl_only': '仅爬取', 'email_only': '仅发送邮件', 'none': '不执行'}
        desc = f"运营商: {int(op_hour):02d}:{int(op_minute):02d} ({range_names.get(op_range, op_range)}), 市场: {int(mkt_hour):02d}:{int(mkt_minute):02d} ({range_names.get(mkt_range, mkt_range)})"
        return jsonify({"status": "success", "message": f"定时任务已启动！{desc}"})
    except Exception as e:
        scheduler_status['running'] = False
        db.save_task_log('scheduler', 'error', str(e))
        logger.exception("定时任务启动失败")
        return jsonify({"status": "error", "message": str(e)})


@app.route('/run-operator-now', methods=['POST'])
def run_operator_now():
    """立即执行运营商爬取任务"""
    try:
        from crawlers import OperatorCrawler
        operator_crawler = OperatorCrawler(app_config)
        data, source_results = operator_crawler.crawl()
        count = len(data) if data else 0
        saved = 0
        if data:
            saved = db.save_operator_data(data)
        for sr in source_results:
            db.save_crawl_log(
                source_key=sr.get('source_key', ''),
                source_name=sr.get('source_name', ''),
                source_url=sr.get('source_url', ''),
                task_type='operator',
                status=sr.get('status', 'unknown'),
                message=sr.get('message', ''),
                found_count=sr.get('found_count', 0),
                saved_count=saved if sr.get('status') == 'success' else 0
            )
        db.save_task_log('operator', 'success', f'获取{count}条,保存{saved}条', count)
        logger.info(f"运营商爬取成功，获取{count}条，保存{saved}条")
        return jsonify({
            "status": "success",
            "message": f"爬取完成！获取{count}条，保存{saved}条（去重后）",
            "count": count,
            "saved": saved,
            "source_results": source_results
        })
    except Exception as e:
        db.save_task_log('operator', 'error', str(e))
        logger.exception("运营商爬取失败")
        return jsonify({"status": "error", "message": f"爬取失败：{str(e)}"})


@app.route('/run-market-now', methods=['POST'])
def run_market_now():
    """立即执行市场爬取任务"""
    try:
        from crawlers import MarketCrawler
        market_crawler = MarketCrawler(app_config)
        data = market_crawler.crawl()
        competitors = data.get('competitors', [])
        hardware = data.get('hardware', [])
        comp_sources = data.get('competitor_sources', [])
        hw_sources = data.get('hardware_sources', [])
        comp_saved = 0
        hw_saved = 0
        if competitors:
            comp_saved = db.save_competitor_data(competitors)
        if hardware:
            hw_saved = db.save_hardware_data(hardware)
        for sr in comp_sources:
            db.save_crawl_log(
                source_key=sr.get('source_key', ''),
                source_name=sr.get('source_name', ''),
                source_url=sr.get('source_url', ''),
                task_type='competitor',
                status=sr.get('status', 'unknown'),
                message=sr.get('message', ''),
                found_count=sr.get('found_count', 0),
                saved_count=comp_saved if sr.get('status') == 'success' else 0
            )
        for sr in hw_sources:
            db.save_crawl_log(
                source_key=sr.get('source_key', ''),
                source_name=sr.get('source_name', ''),
                source_url=sr.get('source_url', ''),
                task_type='hardware',
                status=sr.get('status', 'unknown'),
                message=sr.get('message', ''),
                found_count=sr.get('found_count', 0),
                saved_count=hw_saved if sr.get('status') == 'success' else 0
            )
        total = len(competitors) + len(hardware)
        total_saved = comp_saved + hw_saved
        db.save_task_log('market', 'success', f'获取{total}条(友商{len(competitors)}+硬件{len(hardware)}),保存{total_saved}条', total)
        logger.info(f"市场爬取成功，获取{total}条，保存{total_saved}条")
        return jsonify({
            "status": "success",
            "message": f"爬取完成！友商{len(competitors)}条(保存{comp_saved}条)+硬件{len(hardware)}条(保存{hw_saved}条)",
            "count": total,
            "saved": total_saved,
            "competitor_sources": comp_sources,
            "hardware_sources": hw_sources
        })
    except Exception as e:
        db.save_task_log('market', 'error', str(e))
        logger.exception("市场爬取失败")
        return jsonify({"status": "error", "message": f"爬取失败：{str(e)}"})


# ============ 配置管理API ============
@app.route('/api/config', methods=['GET'])
def get_config():
    try:
        return jsonify({"status": "success", "data": app_config.to_dict()})
    except Exception as e:
        logger.exception("获取配置失败")
        return jsonify({"status": "error", "message": str(e)})


@app.route('/api/config', methods=['POST'])
def update_config():
    try:
        data = request.get_json()
        app_config.update_from_dict(data)
        logger.info("配置更新成功")
        return jsonify({"status": "success", "message": "配置更新成功"})
    except Exception as e:
        logger.exception("配置更新失败")
        return jsonify({"status": "error", "message": str(e)})


# ============ 历史记录API ============
@app.route('/api/history/operator', methods=['GET'])
def get_operator_history():
    try:
        limit = request.args.get('limit', 50, type=int)
        offset = request.args.get('offset', 0, type=int)
        data = db.get_operator_history(limit, offset)
        return jsonify({"status": "success", "data": data})
    except Exception as e:
        logger.exception("获取运营商历史失败")
        return jsonify({"status": "error", "message": str(e)})


@app.route('/api/history/competitor', methods=['GET'])
def get_competitor_history():
    try:
        limit = request.args.get('limit', 50, type=int)
        offset = request.args.get('offset', 0, type=int)
        data = db.get_competitor_history(limit, offset)
        return jsonify({"status": "success", "data": data})
    except Exception as e:
        logger.exception("获取友商历史失败")
        return jsonify({"status": "error", "message": str(e)})


@app.route('/api/history/hardware', methods=['GET'])
def get_hardware_history():
    try:
        limit = request.args.get('limit', 50, type=int)
        offset = request.args.get('offset', 0, type=int)
        data = db.get_hardware_history(limit, offset)
        return jsonify({"status": "success", "data": data})
    except Exception as e:
        logger.exception("获取硬件历史失败")
        return jsonify({"status": "error", "message": str(e)})


@app.route('/api/statistics', methods=['GET'])
def get_statistics():
    try:
        stats = db.get_statistics()
        return jsonify({"status": "success", "data": stats})
    except Exception as e:
        logger.exception("获取统计失败")
        return jsonify({"status": "error", "message": str(e)})


@app.route('/api/task-logs', methods=['GET'])
def get_task_logs():
    try:
        limit = request.args.get('limit', 20, type=int)
        data = db.get_task_logs(limit)
        return jsonify({"status": "success", "data": data})
    except Exception as e:
        logger.exception("获取任务日志失败")
        return jsonify({"status": "error", "message": str(e)})


@app.route('/api/logs', methods=['GET'])
def get_logs():
    try:
        lines = request.args.get('lines', 100, type=int)
        log_file = os.path.join(BASE_DIR, 'app.log')
        if not os.path.exists(log_file):
            return jsonify({"status": "success", "data": []})
        with open(log_file, 'r', encoding='utf-8', errors='ignore') as f:
            all_lines = f.readlines()
        data = all_lines[-lines:]
        return jsonify({"status": "success", "data": [l.strip() for l in data]})
    except Exception as e:
        logger.exception("获取日志失败")
        return jsonify({"status": "error", "message": str(e)})


@app.route('/api/crawl-logs', methods=['GET'])
def get_crawl_logs():
    try:
        task_type = request.args.get('type', '')
        limit = request.args.get('limit', 50, type=int)
        data = db.get_crawl_logs(task_type, limit)
        return jsonify({"status": "success", "data": data})
    except Exception as e:
        logger.exception("获取爬取记录失败")
        return jsonify({"status": "error", "message": str(e)})


@app.route('/api/scheduler-status', methods=['GET'])
def get_scheduler_status():
    return jsonify({"status": "success", "data": scheduler_status})


# ============ 数据操作API ============
@app.route('/api/operator/<int:record_id>', methods=['DELETE'])
def delete_operator(record_id):
    try:
        db.delete_operator(record_id)
        return jsonify({"status": "success", "message": "删除成功"})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})


@app.route('/api/competitor/<int:record_id>', methods=['DELETE'])
def delete_competitor(record_id):
    try:
        db.delete_competitor(record_id)
        return jsonify({"status": "success", "message": "删除成功"})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})


@app.route('/api/hardware/<int:record_id>', methods=['DELETE'])
def delete_hardware(record_id):
    try:
        db.delete_hardware(record_id)
        return jsonify({"status": "success", "message": "删除成功"})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})


@app.route('/send-email/<etype>', methods=['POST'])
def send_email(etype):
    try:
        from email_sender import EmailSender
        sender = EmailSender(app_config)
        if etype == 'operator':
            data = db.get_operator_history(50)
            sender.send_operator_report(data)
        elif etype == 'competitor':
            data = db.get_competitor_history(50)
            sender.send_market_report({'competitors': data, 'hardware': []})
        elif etype == 'hardware':
            data = db.get_hardware_history(50)
            sender.send_market_report({'competitors': [], 'hardware': data})
        elif etype == 'market':
            comp = db.get_competitor_history(50)
            hw = db.get_hardware_history(50)
            sender.send_market_report({'competitors': comp, 'hardware': hw})
        else:
            return jsonify({"status": "error", "message": "未知邮件类型"})
        db.save_task_log(f'email_{etype}', 'success', f'{etype}邮件发送成功')
        return jsonify({"status": "success", "message": "邮件发送成功"})
    except Exception as e:
        db.save_task_log(f'email_{etype}', 'error', str(e))
        logger.exception(f"{etype}邮件发送失败")
        return jsonify({"status": "error", "message": str(e)})


# ============ 搜索筛选API ============
@app.route('/api/search/<data_type>', methods=['GET'])
def search_data(data_type):
    try:
        keyword = request.args.get('keyword', '')
        date_from = request.args.get('date_from', '')
        date_to = request.args.get('date_to', '')
        limit = request.args.get('limit', 50, type=int)
        offset = request.args.get('offset', 0, type=int)
        if data_type == 'operator':
            data = db.search_operator(keyword, date_from, date_to, limit, offset)
        elif data_type == 'competitor':
            data = db.search_competitor(keyword, date_from, date_to, limit, offset)
        elif data_type == 'hardware':
            data = db.search_hardware(keyword, date_from, date_to, limit, offset)
        else:
            return jsonify({"status": "error", "message": "未知的数据类型"})
        return jsonify({"status": "success", "data": data})
    except Exception as e:
        logger.exception(f"搜索{data_type}数据失败")
        return jsonify({"status": "error", "message": str(e)})


# ============ 数据趋势API ============
@app.route('/api/trends', methods=['GET'])
def get_trends():
    try:
        days = request.args.get('days', 7, type=int)
        trends = db.get_trend_data(days)
        return jsonify({"status": "success", "data": trends})
    except Exception as e:
        logger.exception("获取趋势数据失败")
        return jsonify({"status": "error", "message": str(e)})


# ============ 系统状态API ============
@app.route('/api/system-status', methods=['GET'])
def get_system_status():
    try:
        stats = db.get_statistics()
        status = {
            'scheduler_running': scheduler_status['running'],
            'scheduler_started_at': scheduler_status.get('started_at'),
            'total_records': stats['operator_count'] + stats['competitor_count'] + stats['hardware_count'],
            'operator_count': stats['operator_count'],
            'competitor_count': stats['competitor_count'],
            'hardware_count': stats['hardware_count'],
            'today_tasks': stats['today_tasks'],
            'total_tasks': stats['task_count'],
            'uptime': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'version': '3.0'
        }
        return jsonify({"status": "success", "data": status})
    except Exception as e:
        logger.exception("获取系统状态失败")
        return jsonify({"status": "error", "message": str(e)})


# ============ 数据导出API ============
@app.route('/api/export/<data_type>', methods=['GET'])
def export_data(data_type):
    try:
        if data_type == 'operator':
            data = db.get_operator_history(10000)
            headers = ['id', 'title', 'url', 'date', 'source', 'keyword', 'summary', 'crawled_at']
            filename = f'运营商数据_{datetime.now().strftime("%Y%m%d")}.csv'
        elif data_type == 'competitor':
            data = db.get_competitor_history(10000)
            headers = ['id', 'name', 'title', 'url', 'date', 'summary', 'crawled_at']
            filename = f'友商数据_{datetime.now().strftime("%Y%m%d")}.csv'
        elif data_type == 'hardware':
            data = db.get_hardware_history(10000)
            headers = ['id', 'name', 'category', 'price', 'trend', 'url', 'summary', 'crawled_at']
            filename = f'硬件数据_{datetime.now().strftime("%Y%m%d")}.csv'
        else:
            return jsonify({"status": "error", "message": "未知的数据类型"})

        output = io.StringIO()
        writer = csv.DictWriter(output, fieldnames=headers, extrasaction='ignore')
        writer.writeheader()
        writer.writerows(data)
        return Response(
            output.getvalue(),
            mimetype='text/csv',
            headers={'Content-Disposition': f'attachment; filename={filename}'}
        )
    except Exception as e:
        logger.exception(f"导出{data_type}数据失败")
        return jsonify({"status": "error", "message": str(e)})


@app.route('/api/cleanup', methods=['POST'])
def cleanup_data():
    try:
        days = request.args.get('days', 30, type=int)
        db.delete_old_data(days)
        logger.info(f"已清理{days}天前的旧数据")
        return jsonify({"status": "success", "message": f"已清理{days}天前的旧数据"})
    except Exception as e:
        logger.exception("数据清理失败")
        return jsonify({"status": "error", "message": str(e)})


@app.route('/api/seed-demo', methods=['POST'])
def seed_demo_data():
    try:
        db.seed_demo_data()
        return jsonify({"status": "success", "message": "样例数据插入成功"})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})


@app.route('/export/<etype>')
def export_data_simple(etype):
    try:
        if etype == 'operator':
            data = db.get_operator_history(500)
            headers = ['标题', 'URL', '日期', '来源', '关键词', '摘要']
            rows = [[d.get('title', ''), d.get('url', ''), d.get('date', ''),
                     d.get('source', ''), d.get('keyword', ''), d.get('summary', '')] for d in data]
        elif etype == 'competitor':
            data = db.get_competitor_history(500)
            headers = ['友商', '标题', 'URL', '日期', '摘要']
            rows = [[d.get('name', ''), d.get('title', ''), d.get('url', ''),
                     d.get('date', ''), d.get('summary', '')] for d in data]
        elif etype == 'hardware':
            data = db.get_hardware_history(500)
            headers = ['名称', '类别', '价格', '趋势', 'URL', '摘要']
            rows = [[d.get('name', ''), d.get('category', ''), d.get('price', ''),
                     d.get('trend', ''), d.get('url', ''), d.get('summary', '')] for d in data]
        else:
            return jsonify({"status": "error", "message": "未知导出类型"})

        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(headers)
        writer.writerows(rows)
        output.seek(0)
        return Response(
            output.getvalue(),
            mimetype='text/csv',
            headers={'Content-Disposition': f'attachment; filename={etype}_{datetime.now().strftime("%Y%m%d")}.csv'}
        )
    except Exception as e:
        logger.exception(f"导出{etype}数据失败")
        return jsonify({"status": "error", "message": str(e)})


# ============ AI助手API ============
@app.route('/api/ai/models', methods=['GET'])
def get_ai_models():
    """获取可用的AI模型列表（分组结构）"""
    try:
        from ai_service import AI_MODELS, PROVIDER_API_CONFIG, load_ai_config, _get_provider_credentials
        config = load_ai_config()
        current_provider = config.get('currentProvider', 'fusion')
        current_model = config.get('currentModel', 'deepseek-chat')
        groups = []
        # 融合通道组
        fusion_models = AI_MODELS.get('fusion', {})
        fusion_creds = _get_provider_credentials('fusion')
        fusion_has_key = bool(fusion_creds[1])
        groups.append({
            'group': 'fusion',
            'label': '融合通道（通过融合API转发，一个Key用多模型）',
            'has_key': fusion_has_key,
            'models': [{'provider': 'fusion', 'model': m, 'label': m, 'has_key': fusion_has_key} for m in fusion_models.get('models', [])]
        })
        # 单一模型直连组
        direct_items = []
        for key, cfg in AI_MODELS.items():
            if key == 'fusion':
                continue
            creds = _get_provider_credentials(key)
            has_key = bool(creds[1])
            for m in cfg.get('models', []):
                direct_items.append({
                    'provider': key,
                    'model': m,
                    'label': f"{cfg['name']} / {m}",
                    'has_key': has_key
                })
        groups.append({
            'group': 'direct',
            'label': '单一模型直连（需配置对应API Key）',
            'models': direct_items
        })
        return jsonify({
            'status': 'success',
            'data': {
                'groups': groups,
                'currentProvider': current_provider,
                'currentModel': current_model
            }
        })
    except Exception as e:
        logger.exception("获取AI模型列表失败")
        return jsonify({"status": "error", "message": str(e)})


@app.route('/api/ai/config', methods=['GET'])
def get_ai_config():
    """获取AI配置（含各provider的key状态）"""
    try:
        from ai_service import load_ai_config, PROVIDER_API_CONFIG, AI_MODELS
        config = load_ai_config()
        providers_info = {}
        for key, api_cfg in PROVIDER_API_CONFIG.items():
            provider_cfg = config.get('providers', {}).get(key, {})
            has_key = bool(provider_cfg.get('api_key') or api_cfg.get('default_key', ''))
            providers_info[key] = {
                'name': api_cfg['name'],
                'base_url': provider_cfg.get('base_url') or api_cfg['base_url'],
                'has_key': has_key,
                'key_label': api_cfg['key_label'],
                'key_placeholder': api_cfg['key_placeholder'],
                'models': AI_MODELS.get(key, {}).get('models', [])
            }
        return jsonify({
            'status': 'success',
            'data': {
                'currentProvider': config.get('currentProvider', 'fusion'),
                'currentModel': config.get('currentModel', 'deepseek-chat'),
                'providers': providers_info
            }
        })
    except Exception as e:
        logger.exception("获取AI配置失败")
        return jsonify({"status": "error", "message": str(e)})


@app.route('/api/ai/config', methods=['POST'])
def save_ai_config_route():
    """保存AI配置（各provider的Key和当前选择）"""
    try:
        from ai_service import load_ai_config, save_ai_config
        data = request.get_json(force=True, silent=True) or {}
        config = load_ai_config()
        # 更新当前选择
        if 'currentProvider' in data:
            config['currentProvider'] = data['currentProvider']
        if 'currentModel' in data:
            config['currentModel'] = data['currentModel']
        # 更新各provider的凭据
        if 'providers' in data:
            for key, val in data['providers'].items():
                if 'providers' not in config:
                    config['providers'] = {}
                if key not in config['providers']:
                    config['providers'][key] = {}
                if 'api_key' in val:
                    # 空字符串表示清除key，特殊值'__unchanged__'表示不修改
                    if val['api_key'] != '__unchanged__':
                        if val['api_key']:
                            config['providers'][key]['api_key'] = val['api_key']
                        else:
                            config['providers'][key].pop('api_key', None)
                if 'base_url' in val:
                    config['providers'][key]['base_url'] = val['base_url']
        save_ai_config(config)
        logger.info("AI配置保存成功")
        return jsonify({"status": "success", "message": "AI配置保存成功"})
    except Exception as e:
        logger.exception("保存AI配置失败")
        return jsonify({"status": "error", "message": str(e)})


@app.route('/api/ai/chat', methods=['POST'])
def ai_chat():
    """AI对话接口"""
    try:
        from ai_service import call_ai_with_fallback
        data = request.get_json(force=True, silent=True) or {}
        messages = data.get('messages', [])
        provider = data.get('provider')
        model = data.get('model')
        temperature = data.get('temperature', 0.7)
        max_tokens = data.get('max_tokens', 4096)
        if not messages:
            return jsonify({"status": "error", "message": "消息不能为空"}), 400
        result = call_ai_with_fallback(provider=provider, model=model, messages=messages,
                                        temperature=temperature, max_tokens=max_tokens)
        return jsonify({"status": "success", "data": result})
    except Exception as e:
        logger.exception("AI对话失败")
        return jsonify({"status": "error", "message": str(e)})


@app.route('/api/ai/search', methods=['POST'])
def ai_search_route():
    """AI联网搜索接口"""
    try:
        from ai_service import ai_search
        data = request.get_json(force=True, silent=True) or {}
        query = data.get('query', '').strip()
        provider = data.get('provider')
        model = data.get('model')
        if not query:
            return jsonify({"status": "error", "message": "搜索内容不能为空"}), 400
        result = ai_search(query, provider=provider, model=model)
        return jsonify({"status": "success", "data": result})
    except Exception as e:
        logger.exception("AI搜索失败")
        return jsonify({"status": "error", "message": str(e)})


@app.route('/api/ai/export-doc', methods=['POST'])
def ai_export_doc():
    """导出AI对话内容为文档"""
    try:
        from docx import Document
        from docx.shared import RGBColor, Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        data = request.get_json(force=True, silent=True) or {}
        title = data.get('title', 'AI对话记录')
        content = data.get('content', '')
        fmt = data.get('format', 'docx')
        timestamp = data.get('timestamp', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        if not content:
            return jsonify({"status": "error", "message": "内容为空"}), 400
        if fmt == 'txt':
            txt_content = title + '\n生成时间：' + timestamp + '\n\n' + content
            output = io.BytesIO()
            output.write(txt_content.encode('utf-8'))
            output.seek(0)
            return Response(
                output.getvalue(),
                mimetype='text/plain; charset=utf-8',
                headers={'Content-Disposition': f'attachment; filename={title}.txt'}
            )
        else:
            doc = Document()
            # 标题居中
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
                run.font.size = Pt(22)
            # 生成时间居中灰色
            time_para = doc.add_paragraph()
            time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            time_run = time_para.add_run(f'生成时间：{timestamp}')
            time_run.font.size = Pt(10)
            time_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            doc.add_paragraph('')
            # 解析内容
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                # 对话角色标签
                if line.startswith('【用户】'):
                    p = doc.add_paragraph()
                    run = p.add_run('👤 用户')
                    run.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
                    continue
                elif line.startswith('【AI助手】'):
                    p = doc.add_paragraph()
                    run = p.add_run('🤖 AI助手')
                    run.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0x00, 0xbf, 0xa5)
                    continue
                # Markdown标题
                if line.startswith('# '):
                    h = doc.add_heading(line[2:], level=1)
                    for r in h.runs:
                        r.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
                elif line.startswith('## '):
                    h = doc.add_heading(line[3:], level=2)
                    for r in h.runs:
                        r.font.color.rgb = RGBColor(0x2d, 0x3a, 0xbe)
                elif line.startswith('### '):
                    h = doc.add_heading(line[4:], level=3)
                    for r in h.runs:
                        r.font.color.rgb = RGBColor(0x3f, 0x51, 0xb5)
                elif line.startswith('- ') or line.startswith('• '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif re.match(r'^\d+\.\s', line):
                    text = re.sub(r'^\d+\.\s', '', line)
                    doc.add_paragraph(text, style='List Number')
                else:
                    p = doc.add_paragraph()
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
                        run.font.size = Pt(11)
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return Response(
                output.getvalue(),
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={'Content-Disposition': f'attachment; filename={title}.docx'}
            )
    except Exception as e:
        logger.exception("导出文档失败")
        return jsonify({"status": "error", "message": str(e)})


@app.route('/api/ai/generate-weekly', methods=['POST'])
def ai_generate_weekly():
    """AI生成周报"""
    try:
        from ai_service import call_ai_with_fallback, get_knowledge_context
        data = request.get_json(force=True, silent=True) or {}
        provider = data.get('provider')
        model = data.get('model')
        # 获取本周数据
        op_data = db.get_operator_history(50)
        comp_data = db.get_competitor_history(50)
        hw_data = db.get_hardware_history(50)
        op_summary = "\n".join([f"- {d.get('title','')}({d.get('date','')})" for d in op_data[:20]])
        comp_summary = "\n".join([f"- [{d.get('name','')}] {d.get('title','')}({d.get('date','')})" for d in comp_data[:20]])
        hw_summary = "\n".join([f"- {d.get('name','')} 价格:{d.get('price','')} 趋势:{d.get('trend','')}" for d in hw_data[:20]])
        knowledge = get_knowledge_context(3000)
        system_msg = "你是一个专业的周报撰写助手。请根据提供的数据，撰写一份结构清晰、内容专业的周报。周报应包含：本周概要、运营商招标动态、友商动态、硬件市场行情、下周建议等部分。"
        if knowledge:
            system_msg += f"\n\n参考知识库内容：\n{knowledge}"
        user_msg = f"请根据以下数据生成本周工作周报：\n\n【运营商招标动态】\n{op_summary or '暂无数据'}\n\n【友商动态】\n{comp_summary or '暂无数据'}\n\n【硬件市场行情】\n{hw_summary or '暂无数据'}"
        messages = [
            {'role': 'system', 'content': system_msg},
            {'role': 'user', 'content': user_msg}
        ]
        result = call_ai_with_fallback(provider=provider, model=model, messages=messages, temperature=0.7, max_tokens=4096)
        return jsonify({"status": "success", "data": result})
    except Exception as e:
        logger.exception("AI生成周报失败")
        return jsonify({"status": "error", "message": str(e)})


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5001))
    print(f"""
╔══════════════════════════════════════════╗
║     新城小米虾 + 密信本 集成系统            ║
║     http://localhost:{port}                ║
╚══════════════════════════════════════════╝
    """)
    app.run(host='0.0.0.0', port=port, debug=False, use_reloader=False)
