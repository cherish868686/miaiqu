"""资源池配置 - 文档导出模块"""
import io
import json
from datetime import datetime


def export_to_excel(config, name="资源池配置"):
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = name[:31]
    hfont = Font(bold=True, size=14, color="FFFFFF")
    hfill = PatternFill(start_color="1A237E", end_color="1A237E", fill_type="solid")
    sfont = Font(bold=True, size=12, color="1A237E")
    sfill = PatternFill(start_color="E8EAF6", end_color="E8EAF6", fill_type="solid")
    nfont = Font(size=11)
    mfmt = '#,##0'
    tb = Border(left=Side(style='thin'), right=Side(style='thin'),
                top=Side(style='thin'), bottom=Side(style='thin'))
    row = 1
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
    c = ws.cell(row=row, column=1, value=name + " 配置清单")
    c.font = hfont
    c.fill = hfill
    c.alignment = Alignment(horizontal='center', vertical='center')
    for col in range(1, 7):
        ws.cell(row=row, column=col).fill = hfill
        ws.cell(row=row, column=col).border = tb
    row += 2

    # GPU选项
    gpu = config.get("gpu_options", {})
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
    ws.cell(row=row, column=1, value="一、GPU设备选项").font = sfont
    ws.cell(row=row, column=1).fill = sfill
    for col in range(1, 7):
        ws.cell(row=row, column=col).fill = sfill
        ws.cell(row=row, column=col).border = tb
    row += 1
    gpu_labels = [("support_kvm", "是否支持KVM"), ("strong_management", "是否需要强纳管"),
                  ("project_control", "是否项目极致可控"), ("extreme_cost", "是否为极致成本")]
    for key, label in gpu_labels:
        ws.cell(row=row, column=1, value=label).font = nfont
        ws.cell(row=row, column=1).border = tb
        ws.cell(row=row, column=2, value="是" if gpu.get(key, False) else "否").font = nfont
        ws.cell(row=row, column=2).border = tb
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=6)
        row += 1
    row += 1

    # 通算设备
    comp = config.get("compute_options", {})
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
    ws.cell(row=row, column=1, value="二、通算设备选项").font = sfont
    ws.cell(row=row, column=1).fill = sfill
    for col in range(1, 7):
        ws.cell(row=row, column=col).fill = sfill
        ws.cell(row=row, column=col).border = tb
    row += 1
    comp_labels = [("special_business", "是否有特殊业务要求"), ("default_config", "是否默认配置"),
                   ("add_devices", "是否增加设备")]
    for key, label in comp_labels:
        ws.cell(row=row, column=1, value=label).font = nfont
        ws.cell(row=row, column=1).border = tb
        ws.cell(row=row, column=2, value="是" if comp.get(key, False) else "否").font = nfont
        ws.cell(row=row, column=2).border = tb
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=6)
        row += 1
    # 设备列表
    devices = comp.get("devices", [])
    if devices:
        headers = ["设备名称", "规格", "数量", "单价(元)", "小计(元)", "自动填充"]
        for ci, h in enumerate(headers):
            c = ws.cell(row=row, column=ci+1, value=h)
            c.font = Font(bold=True, size=11)
            c.fill = PatternFill(start_color="C5CAE9", end_color="C5CAE9", fill_type="solid")
            c.border = tb
        row += 1
        for dev in devices:
            ws.cell(row=row, column=1, value=dev.get("name","")).font = nfont
            ws.cell(row=row, column=1).border = tb
            ws.cell(row=row, column=2, value=dev.get("spec","")).font = nfont
            ws.cell(row=row, column=2).border = tb
            ws.cell(row=row, column=3, value=dev.get("quantity",0)).font = nfont
            ws.cell(row=row, column=3).border = tb
            ws.cell(row=row, column=4, value=dev.get("unit_price",0)).font = nfont
            ws.cell(row=row, column=4).number_format = mfmt
            ws.cell(row=row, column=4).border = tb
            subtotal = dev.get("quantity",0) * dev.get("unit_price",0)
            ws.cell(row=row, column=5, value=subtotal).font = nfont
            ws.cell(row=row, column=5).number_format = mfmt
            ws.cell(row=row, column=5).border = tb
            ws.cell(row=row, column=6, value="是" if dev.get("auto_fill") else "否").font = nfont
            ws.cell(row=row, column=6).border = tb
            row += 1
    row += 1

    # 网络设备
    net = config.get("network_options", {})
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
    ws.cell(row=row, column=1, value="三、网络设备选项").font = sfont
    ws.cell(row=row, column=1).fill = sfill
    for col in range(1, 7):
        ws.cell(row=row, column=col).fill = sfill
        ws.cell(row=row, column=col).border = tb
    row += 1
    net_labels = [("strong_management", "是否需求强纳管"),
                  ("inband_management_switch", "是否需求带内管理交换机"),
                  ("management_convergence_switch", "是否需求管理汇聚交换机"),
                  ("compute_convergence_switch", "是否需求算力汇聚交换机"),
                  ("edge_exit_switch", "是否需求边缘出口交换机"),
                  ("dedicated_line_switch", "是否需求专线交换机")]
    for key, label in net_labels:
        ws.cell(row=row, column=1, value=label).font = nfont
        ws.cell(row=row, column=1).border = tb
        ws.cell(row=row, column=2, value="是" if net.get(key, False) else "否").font = nfont
        ws.cell(row=row, column=2).border = tb
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=6)
        row += 1
    net_devices = net.get("devices", [])
    if net_devices:
        headers = ["设备名称", "规格", "数量", "单价(元)", "小计(元)", "自动填充"]
        for ci, h in enumerate(headers):
            c = ws.cell(row=row, column=ci+1, value=h)
            c.font = Font(bold=True, size=11)
            c.fill = PatternFill(start_color="C5CAE9", end_color="C5CAE9", fill_type="solid")
            c.border = tb
        row += 1
        for dev in net_devices:
            ws.cell(row=row, column=1, value=dev.get("name","")).font = nfont
            ws.cell(row=row, column=1).border = tb
            ws.cell(row=row, column=2, value=dev.get("spec","")).font = nfont
            ws.cell(row=row, column=2).border = tb
            ws.cell(row=row, column=3, value=dev.get("quantity",0)).font = nfont
            ws.cell(row=row, column=3).border = tb
            ws.cell(row=row, column=4, value=dev.get("unit_price",0)).font = nfont
            ws.cell(row=row, column=4).number_format = mfmt
            ws.cell(row=row, column=4).border = tb
            subtotal = dev.get("quantity",0) * dev.get("unit_price",0)
            ws.cell(row=row, column=5, value=subtotal).font = nfont
            ws.cell(row=row, column=5).number_format = mfmt
            ws.cell(row=row, column=5).border = tb
            ws.cell(row=row, column=6, value="是" if dev.get("auto_fill") else "否").font = nfont
            ws.cell(row=row, column=6).border = tb
            row += 1
    row += 1

    # 成本汇总
    cost = config.get("cost_options", {})
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
    ws.cell(row=row, column=1, value="四、成本汇总").font = sfont
    ws.cell(row=row, column=1).fill = sfill
    for col in range(1, 7):
        ws.cell(row=row, column=col).fill = sfill
        ws.cell(row=row, column=col).border = tb
    row += 1
    # 计算硬件总成本
    hw_total = 0
    for dev in comp.get("devices", []):
        hw_total += dev.get("quantity", 0) * dev.get("unit_price", 0)
    for dev in net.get("devices", []):
        hw_total += dev.get("quantity", 0) * dev.get("unit_price", 0)
    ws.cell(row=row, column=1, value="硬件总成本").font = nfont
    ws.cell(row=row, column=1).border = tb
    ws.cell(row=row, column=2, value=hw_total).font = nfont
    ws.cell(row=row, column=2).number_format = mfmt
    ws.cell(row=row, column=2).border = tb
    ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=6)
    row += 1

    total = hw_total
    if cost.get("add_warranty"):
        warranty = hw_total * cost.get("warranty_rate", 0.15)
        ws.cell(row=row, column=1, value="维保成本").font = nfont
        ws.cell(row=row, column=1).border = tb
        ws.cell(row=row, column=2, value=warranty).font = nfont
        ws.cell(row=row, column=2).number_format = mfmt
        ws.cell(row=row, column=2).border = tb
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=6)
        row += 1
        total += warranty

    if cost.get("add_software"):
        sw_total = sum(item.get("cost", 0) for item in cost.get("software_items", []))
        ws.cell(row=row, column=1, value="软件成本").font = nfont
        ws.cell(row=row, column=1).border = tb
        ws.cell(row=row, column=2, value=sw_total).font = nfont
        ws.cell(row=row, column=2).number_format = mfmt
        ws.cell(row=row, column=2).border = tb
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=6)
        row += 1
        total += sw_total

    if cost.get("add_integration"):
        integration = hw_total * cost.get("integration_rate", 0.08)
        ws.cell(row=row, column=1, value="集成成本").font = nfont
        ws.cell(row=row, column=1).border = tb
        ws.cell(row=row, column=2, value=integration).font = nfont
        ws.cell(row=row, column=2).number_format = mfmt
        ws.cell(row=row, column=2).border = tb
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=6)
        row += 1
        total += integration

    row += 1
    ws.cell(row=row, column=1, value="总成本").font = Font(bold=True, size=13, color="D32F2F")
    ws.cell(row=row, column=1).border = tb
    ws.cell(row=row, column=2, value=total).font = Font(bold=True, size=13, color="D32F2F")
    ws.cell(row=row, column=2).number_format = mfmt
    ws.cell(row=row, column=2).border = tb
    ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=6)

    for col in [1, 2, 3, 4, 5, 6]:
        ws.column_dimensions[chr(64+col)].width = [30, 40, 10, 15, 15, 12][col-1]

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


def export_to_word(config, name="资源池配置"):
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(11)
    title = doc.add_heading(name + " 配置清单", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("生成时间: " + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    doc.add_paragraph()

    # GPU选项
    doc.add_heading("一、GPU设备选项", level=1)
    gpu = config.get("gpu_options", {})
    gpu_labels = [("support_kvm", "是否支持KVM"), ("strong_management", "是否需要强纳管"),
                  ("project_control", "是否项目极致可控"), ("extreme_cost", "是否为极致成本")]
    table = doc.add_table(rows=1, cols=2, style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].text = "选项"
    table.rows[0].cells[1].text = "值"
    for key, label in gpu_labels:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = "是" if gpu.get(key, False) else "否"
    doc.add_paragraph()

    # 通算设备
    doc.add_heading("二、通算设备选项", level=1)
    comp = config.get("compute_options", {})
    comp_labels = [("special_business", "是否有特殊业务要求"), ("default_config", "是否默认配置"),
                   ("add_devices", "是否增加设备")]
    table = doc.add_table(rows=1, cols=2, style='Light Grid Accent 1')
    table.rows[0].cells[0].text = "选项"
    table.rows[0].cells[1].text = "值"
    for key, label in comp_labels:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = "是" if comp.get(key, False) else "否"
    doc.add_paragraph()
    devices = comp.get("devices", [])
    if devices:
        doc.add_heading("通算设备清单", level=2)
        table = doc.add_table(rows=1, cols=5, style='Light Grid Accent 1')
        for ci, h in enumerate(["设备名称", "规格", "数量", "单价(元)", "小计(元)"]):
            table.rows[0].cells[ci].text = h
        for dev in devices:
            row = table.add_row()
            row.cells[0].text = dev.get("name", "")
            row.cells[1].text = dev.get("spec", "")
            row.cells[2].text = str(dev.get("quantity", 0))
            row.cells[3].text = str(dev.get("unit_price", 0))
            row.cells[4].text = str(dev.get("quantity", 0) * dev.get("unit_price", 0))
    doc.add_paragraph()

    # 网络设备
    doc.add_heading("三、网络设备选项", level=1)
    net = config.get("network_options", {})
    net_labels = [("strong_management", "是否需求强纳管"),
                  ("inband_management_switch", "是否需求带内管理交换机"),
                  ("management_convergence_switch", "是否需求管理汇聚交换机"),
                  ("compute_convergence_switch", "是否需求算力汇聚交换机"),
                  ("edge_exit_switch", "是否需求边缘出口交换机"),
                  ("dedicated_line_switch", "是否需求专线交换机")]
    table = doc.add_table(rows=1, cols=2, style='Light Grid Accent 1')
    table.rows[0].cells[0].text = "选项"
    table.rows[0].cells[1].text = "值"
    for key, label in net_labels:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = "是" if net.get(key, False) else "否"
    doc.add_paragraph()
    net_devices = net.get("devices", [])
    if net_devices:
        doc.add_heading("网络设备清单", level=2)
        table = doc.add_table(rows=1, cols=5, style='Light Grid Accent 1')
        for ci, h in enumerate(["设备名称", "规格", "数量", "单价(元)", "小计(元)"]):
            table.rows[0].cells[ci].text = h
        for dev in net_devices:
            row = table.add_row()
            row.cells[0].text = dev.get("name", "")
            row.cells[1].text = dev.get("spec", "")
            row.cells[2].text = str(dev.get("quantity", 0))
            row.cells[3].text = str(dev.get("unit_price", 0))
            row.cells[4].text = str(dev.get("quantity", 0) * dev.get("unit_price", 0))
    doc.add_paragraph()

    # 成本汇总
    doc.add_heading("四、成本汇总", level=1)
    cost = config.get("cost_options", {})
    hw_total = 0
    for dev in comp.get("devices", []):
        hw_total += dev.get("quantity", 0) * dev.get("unit_price", 0)
    for dev in net.get("devices", []):
        hw_total += dev.get("quantity", 0) * dev.get("unit_price", 0)
    table = doc.add_table(rows=1, cols=2, style='Light Grid Accent 1')
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "金额(元)"
    row = table.add_row()
    row.cells[0].text = "硬件总成本"
    row.cells[1].text = str(hw_total)
    total = hw_total
    if cost.get("add_warranty"):
        warranty = hw_total * cost.get("warranty_rate", 0.15)
        row = table.add_row()
        row.cells[0].text = "维保成本"
        row.cells[1].text = str(warranty)
        total += warranty
    if cost.get("add_software"):
        sw_total = sum(item.get("cost", 0) for item in cost.get("software_items", []))
        row = table.add_row()
        row.cells[0].text = "软件成本"
        row.cells[1].text = str(sw_total)
        total += sw_total
    if cost.get("add_integration"):
        integration = hw_total * cost.get("integration_rate", 0.08)
        row = table.add_row()
        row.cells[0].text = "集成成本"
        row.cells[1].text = str(integration)
        total += integration
    row = table.add_row()
    row.cells[0].text = "总成本"
    row.cells[1].text = str(total)
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
